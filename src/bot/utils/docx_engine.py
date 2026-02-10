"""C2. Генератор юридических документов по шаблонам (Docx-Engine).

Генерирует .docx документы на основе ответов пользователя.
Использует python-docx для подстановки данных в шаблоны.
Результат отправляется пользователю и дублируется в облако.

C7. Encrypted Vault — шифрованное хранилище клиентских документов.
"""

import io
import logging
import os
from datetime import datetime, timezone
from pathlib import Path

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path("data/templates")
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

VAULT_DIR = Path("data/vault")
VAULT_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_DIR = Path("data/generated_docs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ═══════════════════════════════════════════════════════════════════════════
#  ШАБЛОНЫ ДОКУМЕНТОВ
# ═══════════════════════════════════════════════════════════════════════════

DOCX_TEMPLATES = {
    "nda": {
        "title": "📝 NDA — Соглашение о неразглашении",
        "fields": ["client_name", "counterparty", "city", "purpose", "duration_months"],
        "questions": [
            ("client_name", "Полное наименование вашей компании (или ваше ФИО):"),
            ("counterparty", "Полное наименование контрагента:"),
            ("city", "Город подписания (или <code>-</code> для Алматы):"),
            ("purpose", "Цель соглашения (или <code>-</code> для стандартной):"),
            ("duration_months", "Срок конфиденциальности в месяцах (или <code>-</code> для 24):"),
        ],
        "price": 0,  # бесплатно
    },
    "power_of_attorney": {
        "title": "📋 Доверенность (общая)",
        "fields": ["principal", "attorney", "scope", "valid_until"],
        "questions": [
            ("principal", "ФИО и данные доверителя (кто выдаёт):"),
            ("attorney", "ФИО и данные поверенного (кому выдаётся):"),
            ("scope", "Полномочия (или <code>-</code> для общих):"),
            ("valid_until", "Срок действия (или <code>-</code> для 1 года):"),
        ],
        "price": 0,
    },
    "claim_letter": {
        "title": "⚠️ Претензия (досудебная)",
        "fields": ["sender", "recipient", "subject", "amount", "deadline_days"],
        "questions": [
            ("sender", "От кого (ваша компания / ФИО):"),
            ("recipient", "Кому (наименование контрагента):"),
            ("subject", "Суть претензии (кратко):"),
            ("amount", "Сумма требования (или <code>-</code> если не денежная):"),
            ("deadline_days", "Срок для ответа в днях (или <code>-</code> для 15):"),
        ],
        "price": 0,
    },
    "employment_contract": {
        "title": "👔 Трудовой договор",
        "fields": ["employer", "employee", "position", "salary", "start_date"],
        "questions": [
            ("employer", "Наименование работодателя:"),
            ("employee", "ФИО работника:"),
            ("position", "Должность:"),
            ("salary", "Размер оклада (тенге/месяц):"),
            ("start_date", "Дата начала (или <code>-</code> для текущей):"),
        ],
        "price": 5000,  # 5000 тенге
    },
    "service_agreement": {
        "title": "🤝 Договор оказания услуг",
        "fields": ["executor", "client", "service_desc", "amount", "deadline"],
        "questions": [
            ("executor", "Наименование исполнителя:"),
            ("client", "Наименование заказчика:"),
            ("service_desc", "Описание услуги:"),
            ("amount", "Стоимость (тенге, или <code>-</code> для «по согласованию»):"),
            ("deadline", "Срок исполнения (или <code>-</code> для 30 дней):"),
        ],
        "price": 5000,
    },
}


async def generate_document_docx(template_id: str, data: dict, user_id: int = 0) -> str | None:
    """Генерирует .docx документ по шаблону и данным.

    Сначала проверяет наличие файла шаблона (.docx).
    Если шаблона нет — генерирует через AI.

    Returns:
        Путь к сгенерированному файлу или None.
    """
    template = DOCX_TEMPLATES.get(template_id)
    if not template:
        logger.warning("Unknown template: %s", template_id)
        return None

    # Проверяем наличие .docx шаблона
    tmpl_path = TEMPLATE_DIR / f"{template_id}_template.docx"
    if tmpl_path.exists():
        return await _fill_docx_template(tmpl_path, data, template_id, user_id)

    # Fallback: генерация через AI
    return await _generate_ai_document(template_id, template, data, user_id)


async def _fill_docx_template(
    tmpl_path: Path, data: dict, template_id: str, user_id: int,
) -> str | None:
    """Заполняет .docx шаблон подстановкой переменных."""
    import asyncio

    def _fill():
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed")
            return None

        doc = Document(str(tmpl_path))

        # Подстановка в параграфах
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = "{" + key.upper() + "}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Подстановка в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = "{" + key.upper() + "}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        # Сохраняем
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename = f"{template_id}_{user_id}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename
        doc.save(str(filepath))
        return str(filepath)

    return await asyncio.to_thread(_fill)


async def _generate_ai_document(
    template_id: str, template: dict, data: dict, user_id: int,
) -> str | None:
    """Генерирует документ через AI, когда шаблона нет."""
    from src.bot.utils.ai_client import get_orchestrator
    import asyncio

    ai = get_orchestrator()

    prompt = (
        f"Сгенерируй полноценный юридический документ: {template['title']}\n\n"
        f"Данные:\n"
    )
    for k, v in data.items():
        prompt += f"  • {k}: {v}\n"
    prompt += (
        "\nТребования:\n"
        "1. Полноценный документ по законодательству РК\n"
        "2. Все обязательные реквизиты\n"
        "3. Нумерация пунктов\n"
        "4. Место для подписей сторон\n"
        "5. Формат: чистый текст (без HTML/Markdown)"
    )

    text = await ai.call_with_fallback(
        prompt,
        "Ты — юрист SOLIS Partners. Составляй профессиональные юридические документы по РК.",
        primary="openai", max_tokens=4096, temperature=0.3,
    )

    # Сохраняем как .docx
    try:
        def _save():
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(template["title"].replace("📝 ", "").replace("📋 ", "")
                                .replace("⚠️ ", "").replace("👔 ", "").replace("🤝 ", ""), level=1)
                for para in text.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
                filename = f"{template_id}_{user_id}_{timestamp}.docx"
                filepath = OUTPUT_DIR / filename
                doc.save(str(filepath))
                return str(filepath)
            except ImportError:
                # python-docx not available — save as .txt
                timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
                filename = f"{template_id}_{user_id}_{timestamp}.txt"
                filepath = OUTPUT_DIR / filename
                filepath.write_text(text, encoding="utf-8")
                return str(filepath)

        return await asyncio.to_thread(_save)
    except Exception as e:
        logger.error("Document save failed: %s", e)
        return None


def get_document_as_bytes(filepath: str) -> io.BytesIO | None:
    """Читает документ в BytesIO для отправки в Telegram."""
    try:
        with open(filepath, "rb") as f:
            buf = io.BytesIO(f.read())
            buf.seek(0)
            return buf
    except Exception as e:
        logger.error("Failed to read document: %s", e)
        return None


# ═══════════════════════════════════════════════════════════════════════════
#  C7: Encrypted Vault
# ═══════════════════════════════════════════════════════════════════════════

_VAULT_KEY: bytes | None = None


def _get_vault_key() -> bytes:
    """Получает или генерирует ключ шифрования."""
    global _VAULT_KEY
    if _VAULT_KEY:
        return _VAULT_KEY

    key_file = VAULT_DIR / ".vault_key"
    if key_file.exists():
        _VAULT_KEY = key_file.read_bytes()
    else:
        try:
            from cryptography.fernet import Fernet
            _VAULT_KEY = Fernet.generate_key()
            key_file.write_bytes(_VAULT_KEY)
        except ImportError:
            # Без cryptography — используем base64 encode
            import base64
            _VAULT_KEY = base64.urlsafe_b64encode(os.urandom(32))
            key_file.write_bytes(_VAULT_KEY)

    return _VAULT_KEY


async def encrypt_and_store(
    data: bytes, filename: str, user_id: int, metadata: dict | None = None,
) -> str:
    """Шифрует данные и сохраняет в vault.

    Returns:
        Путь к зашифрованному файлу.
    """
    import asyncio

    def _store():
        key = _get_vault_key()
        try:
            from cryptography.fernet import Fernet
            f = Fernet(key)
            encrypted = f.encrypt(data)
        except ImportError:
            # Без cryptography — XOR с ключом (базовое шифрование)
            encrypted = bytes(b ^ key[i % len(key)] for i, b in enumerate(data))

        user_dir = VAULT_DIR / str(user_id)
        user_dir.mkdir(exist_ok=True)

        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        safe_name = f"{timestamp}_{filename}"
        filepath = user_dir / safe_name
        filepath.write_bytes(encrypted)

        # Сохраняем метаданные
        meta_path = filepath.with_suffix(filepath.suffix + ".meta")
        import json
        meta = {
            "original_name": filename,
            "user_id": user_id,
            "stored_at": datetime.now(timezone.utc).isoformat(),
            "size": len(data),
            **(metadata or {}),
        }
        meta_path.write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")

        return str(filepath)

    return await asyncio.to_thread(_store)


async def decrypt_and_retrieve(filepath: str) -> bytes | None:
    """Расшифровывает файл из vault."""
    import asyncio

    def _retrieve():
        key = _get_vault_key()
        encrypted = Path(filepath).read_bytes()
        try:
            from cryptography.fernet import Fernet
            f = Fernet(key)
            return f.decrypt(encrypted)
        except ImportError:
            return bytes(b ^ key[i % len(key)] for i, b in enumerate(encrypted))
        except Exception as e:
            logger.error("Decrypt failed: %s", e)
            return None

    return await asyncio.to_thread(_retrieve)


def get_user_vault_files(user_id: int) -> list[dict]:
    """Возвращает список файлов пользователя в vault."""
    user_dir = VAULT_DIR / str(user_id)
    if not user_dir.exists():
        return []

    files = []
    import json
    for meta_file in user_dir.glob("*.meta"):
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
            files.append(meta)
        except Exception:
            pass

    return sorted(files, key=lambda x: x.get("stored_at", ""), reverse=True)
